import platform
import pandas as pd
import re
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
import requests
from bs4 import BeautifulSoup as bs

import docx
import os

import time
import datetime
from calendar import monthrange


pd.set_option('display.max_rows', None)
pd.set_option('display.max_columns', None)


def str_digit2month(month):
    """
    Функция переводит название месяца в его номер.
    """
    month = month.strip().lower()
    if month == '01':
        return 'Январь'
    elif month == '02':
        return 'Январь-февраль'
    elif month == '03':
        return 'Январь-март'
    elif month == '04':
        return 'Январь-апрель'
    elif month == '05':
        return 'Январь-май'
    elif month == '06':
        return 'Январь-июнь'
    elif month == '07':
        return 'Январь-июль'
    elif month == '08':
        return 'Январь-август'
    elif month == '09':
        return 'Январь-сентябрь'
    elif month == '10':
        return 'Январь-октябрь'
    elif month == '11':
        return 'Январь-ноябрь'
    elif month == '12':
        return 'Январь-декабрь'
    else:
        return 'unknown'


def str_month2digit_month(month):
    """
    Функция переводит название месяца в его номер.
    """
    month = month.strip().lower()
    if month == 'январь':
        return '01'
    elif month == 'январь-февраль':
        return '02'
    elif month == 'январь-март':
        return '03'
    elif month == 'январь-апрель':
        return '04'
    elif month == 'январь-май':
        return '05'
    elif month == 'январь-июнь':
        return '06'
    elif month == 'январь-июль':
        return '07'
    elif month == 'январь-август':
        return '08'
    elif month == 'январь-сентябрь':
        return '09'
    elif month == 'январь-октябрь':
        return '10'
    elif month == 'январь-ноябрь':
        return '11'
    elif month == 'январь-декабрь':
        return '12'
    else:
        return 'unknown'


def reformat_date(date: str, year):
    """
    Функция переформатирует даты
    """
    date = date.strip()
    flag = True if ((year % 4 == 0 and year % 100 != 0) or (year % 400 == 0)) else False
    if date == 'Январь':
        date = '31 january'
    elif date == 'Январь-февраль' and flag:
        date = '29 february'
    elif date == 'Январь-февраль':
        date = '28 february'
    elif date == 'I квартал':
        date = '31 march'
    elif date == 'Январь-апрель':
        date = '30 April'
    elif date == 'Январь-май':
        date = '31 may'
    elif date == 'I полугодие':
        date = '30 june'
    elif date == 'Январь-июль':
        date = '31 july'
    elif date == 'Январь-август':
        date = '31 august'
    elif date == 'Январь-сентябрь':
        date = '30 september'
    elif date == 'Январь-октябрь':
        date = '31 october'
    elif date == 'Январь-ноябрь':
        date = '30 november'
    elif date == 'Год' or date == 'Год1)':
        date = '31 december'
    return date


def reformate_quarterly_date(date):
    if date == 'Январь-март':
        date = 'I квартал'
    elif date == 'Январь-июнь':
        date = 'I полугодие'
    elif date == 'Январь-декабрь':
        date = 'Год'
    return date


def pars_year_by_months(year):
    """
    Функция для получения ссылок на документы по месяцам.
    Для инвестиций реализовано возвращение названия последнего доступного месяца в конкретном году
    и ссылки на соответствующий раздел.
    """
    header = {
        'user-agent': 'Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:86.0) Gecko/20100101 Firefox/86.0'
    }

    url = f'https://rosstat.gov.ru/storage/mediabank/Doklad_{year}.htm'
    session = requests.Session()
    retry = Retry(connect=3, backoff_factor=0.5)
    adapter = HTTPAdapter(max_retries=retry)
    session.mount('https://', adapter)

    response = session.get(url, headers=header)
    soup = bs(response.content, "html.parser")

    links_1 = pd.DataFrame()
    for i in range(0, len(soup.find('table').find_all('tr')[1].find_all('tr')), 2):
        month_name = soup.find('table').find_all('tr')[1].find_all('tr')[i].find_all('td')[0].text
        month_name = month_name.replace('\n', '')
        if month_name.split()[-1].lower() == 'год':
            month_name = 'Январь-декабрь'
        dok_link = soup.find('table').find_all('tr')[1].find_all('tr')[i].find_all('td')[1].find_all('a')[0].get('href')
        if dok_link[:4] != 'http':
            dok_link = 'https://rosstat.gov.ru' + dok_link
        pril_link = soup.find('table').find_all('tr')[1].find_all('tr')[i + 1].find_all('td')[0].find_all('a')[0].get(
            'href')
        if pril_link[:4] != 'http':
            pril_link = 'https://rosstat.gov.ru' + pril_link
        links_1 = links_1._append([[month_name, dok_link, pril_link]])

    links_1.columns = ['Месяц', 'Ссылка', 'Дополнительная ссылка']
    links_1 = links_1.iloc[::-1].reset_index(drop=True)

    return links_1


def download_document(year, month, url):
    '''
    Функция скачивает документ с данными по зарплатам за конкретный месяц.
    year - год в формате ХХХХ.
    month - полное название месяца на русском языке.
    url - ссылка на документ.
    Первые две переменные необходимы для назначения имени скачиваемому файлу.
    Возвращает путь к сохранённому файлу.
    '''

    indicator = 'Розничная торговля'
    doc_link = ''

    header = {
        'user-agent': 'Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:86.0) Gecko/20100101 Firefox/86.0'
    }
    month = str_month2digit_month(month)
    session = requests.Session()
    retry = Retry(connect=3, backoff_factor=0.5)
    adapter = HTTPAdapter(max_retries=retry)
    session.mount('https://', adapter)

    response = session.get(url, headers=header)
    soup = bs(response.content, "html.parser")

    for link in soup.find_all('a'):
        branch_name = link.text
        branch_name = branch_name.replace('\n', '').replace('\r', '').strip()
        branch_name = re.sub(' +', ' ', branch_name)
        if branch_name == indicator:
            doc_link = link.get('href')
            break

    if len(doc_link) == 0:
        print(f'NO DOCUMENTS {year}_{month}: {indicator}')
    else:
        link_to_download = doc_link
        dok_name_to_download = f'{year}_{month}-2-4-0.doc'  # 2024_02-2-4-0.doc
        folder = os.getcwd()
        folder = os.path.join(folder, 'word_data', dok_name_to_download)

        response = session.get(link_to_download, headers=header)

        if response.status_code == 200:
            with open(folder, 'wb') as f:
                f.write(response.content)
            print(f'Document {year}_{month} was downloaded.')
        else:
            print('FAILED:', link_to_download)

        return folder


def check_last_month_in_table(links_list, year, xlsx_path='rez_file_Y_v2.xlsx'):
    data_xlsx = pd.read_excel(xlsx_path)
    data_xlsx = data_xlsx[['Целевой показатель', 'Розничный товарооборот']]
    month = data_xlsx.dropna(subset=['Розничный товарооборот']).iloc[-1]['Целевой показатель']
    if year == month.year + 1 and month.month in [12]:
        return links_list, 0
    elif month.month in [12] and year == month.year:
        return links_list[12:], 0
    else:
        month = str_digit2month(str(month).split('-')[1])
        idx = links_list[links_list['Месяц'] == month].index[0]
        return links_list.loc[idx + 1:], idx + 1


def get_os_type():
    os_type = platform.system()
    if os_type == "Windows":
        return "Windows"
    elif os_type == "Darwin" or os_type == "Linux":
        return "Unix"
    else:
        return "Unknown"


def doc_to_docx(path: str):
    """
    Функция конвертирует документ формата .doc в формат .docx
    doc_path - абсолютный путь к документу
    """
    exist_system = get_os_type()
    if exist_system == 'Unix':
        import doc2docx
        doc2docx.convert(path)

    elif exist_system == 'Windows':
        from win32com import client as wc
        w = wc.Dispatch('Word.Application')
        # Or use the following method to start a separate process:
        # w = wc.DispatchEx('Word.Application')
        doc = w.Documents.Open(path)
        doc.SaveAs(path + 'x', 16)
        doc.Close()
        w.Quit()
        print(f'Document {path} was converted to docx-format.')

    return path + 'x'


def parse_docx_document(path, year, month):
    '''
    Функция осуществляет парсинг документа.
    path - путь к документу (обязательно в формате .docx)
    year - текущий год
    '''
    try:
        doc = docx.Document(path)
    except:
        print('parse_docx_document: It is not word document')
        return 0, 0, 0
    data_table = [[] for _ in range(len(doc.tables[0].rows))]
    for i, row in enumerate(doc.tables[0].rows):
        for cell in row.cells:
            data_table[i].append(cell.text)

    data_table = pd.DataFrame(data_table)
    comment = data_table.iloc[-1, 0]
    if month == 'Январь':
        data_table = data_table[data_table.iloc[:, 0].str.contains(f' Январь|Год')]
        data_table = data_table.iloc[-2:]
        data_table.iloc[:, 0] = data_table.iloc[:, 0].apply(lambda x: reformat_date(x, year))
        data_table.iloc[-2, 0] = pd.to_datetime(data_table.iloc[-2, 0] + str(year - 1))

    else:
        data_table = data_table[data_table.iloc[:, 0].str.contains(f'{reformate_quarterly_date(month)}')]
        data_table = data_table.iloc[-1:]
        data_table.iloc[:, 0] = data_table.iloc[:, 0].apply(lambda x: reformat_date(x, year))
    data_table.iloc[-1, 0] = pd.to_datetime(data_table.iloc[-1, 0] + str(year))
    data_table = data_table[[0, 3, 5]]

    return data_table, comment


def create_new_date(last_date_in_file_year, last_date_in_file_month):
    now = datetime.datetime.now()
    lst_date = []
    _, last_day = monthrange(now.year, now.month)
    last_date = datetime.datetime.strptime(f"{now.year}-{now.month}-{last_day}", "%Y-%m-%d").date()

    for i in range((last_date.year - last_date_in_file_year) * 12 + last_date.month - last_date_in_file_month - 1):
        if last_date.month - 1 != 0:
            _, last_day = monthrange(last_date.year, last_date.month - 1)
            last_date = datetime.datetime.strptime(f"{last_date.year}-{last_date.month - 1}-{last_day}", "%Y-%m-%d").date()
        else:
            _, last_day = monthrange(last_date.year - 1, 12)
            last_date = datetime.datetime.strptime(f"{last_date.year - 1}-{12}-{last_day}", "%Y-%m-%d").date()
        lst_date.append(last_date)
    return sorted(lst_date)


def append_date_rez_file_Y(xlsx_path='rez_file_Y_v2.xlsx'):
    """
        Функция осуществляет дабавление месяцев, если их нет в файле.
    """
    data_xlsx = pd.read_excel(xlsx_path)
    year = pd.to_datetime(pd.read_excel('rez_file_Y_v2.xlsx')['Целевой показатель'].iloc[-1]).year
    month = pd.to_datetime(pd.read_excel('rez_file_Y_v2.xlsx')['Целевой показатель'].iloc[-1]).month
    date_lst = create_new_date(year, month)
    for date in date_lst:
        new_string = {'Целевой показатель': [date]}
        new_string.update({c: [None] for c in data_xlsx.columns[1:]})
        new_string = pd.DataFrame(new_string)
        if not data_xlsx.empty and not new_string.empty:
            data_xlsx = pd.concat([data_xlsx, new_string])
    data_xlsx.to_excel(xlsx_path, index=False)


def update_rez_file_y(data, kvartal_data, xlsx_path='rez_file_Y_v2.xlsx'):
    """
        Функция осуществляет обновление файла со всеми данными rez_file_Y_v2.xlsx
    """
    data_xlsx = pd.read_excel(xlsx_path)
    if list(data.keys())[-1] not in list(data_xlsx['Целевой показатель']):
        append_date_rez_file_Y()
        data_xlsx = pd.read_excel(xlsx_path)
    name_1 = 'Розничный товарооборот'
    name_2 = 'Розничный товарооборот, темп роста, % г/г'
    for j in data:
        data_xlsx.loc[data_xlsx['Целевой показатель'] == j, name_1] = data[j]
    if len(kvartal_data) != 0:
        for c in kvartal_data:
            data_xlsx.loc[data_xlsx['Целевой показатель'] == c, name_2] = kvartal_data[c]

    data_xlsx.to_excel(xlsx_path, index=False)
    print(f'rez_file_Y_v2.xlsx was apdated')


def main():
    """
    Основная функция. Выполняет проверку данных на полноту. Скачивет недостающие
    данные и дополняет ими файл с данными.
    """
    now = datetime.datetime.now().year
    last_year_in_table = pd.to_datetime(pd.read_excel('rez_file_Y_v2.xlsx').dropna(subset=['Розничный товарооборот']).iloc[
                                            -1]['Целевой показатель']).year
    new_data = {}
    kvartal_data = {}
    if now - last_year_in_table < 1:
        years = [now]
    else:
        years = []
        for y in range(last_year_in_table, now + 1):
            years.append(y)
    for year in years:
        time.sleep(15)
        links_data = pars_year_by_months(year)
        links_data, idx = check_last_month_in_table(links_data, year)
        if links_data.empty:
            continue
        else:
            print('Ссылки получены')
            for month in links_data['Месяц']:
                # Скачиваем файл и экспортируем его в докс
                URL = list(links_data.iloc[links_data[links_data['Месяц'] == month].index - idx]['Ссылка'])[0]
                print(URL)
                time.sleep(15)
                path_to_docfile = download_document(year, month, URL)
                print(path_to_docfile)
                time.sleep(15)
                path = doc_to_docx(path_to_docfile)
                data, comm = parse_docx_document(path, year=year, month=month)

                if month == 'Январь':
                    last_date, last_value_1, last_value_2 = data.iloc[0, 0], data.iloc[0, 1], data.iloc[0, 2]
                    date, value_1, value_2 = data.iloc[1, 0], data.iloc[1, 1], data.iloc[1, 2]
                    kvartal_data[last_date] = float(last_value_2.replace(',', '.').replace(' ', ''))
                    new_data[last_date] = float(last_value_1.replace(',', '.').replace(' ', ''))
                    new_data[date] = float(value_1.replace(',', '.').replace(' ', ''))
                else:
                    date, value_1, value_2 = data.iloc[0, 0], data.iloc[0, 1], data.iloc[0, 2]
                    if month in ['Январь-март', 'Январь-июнь', 'Январь-сентябрь', 'Январь-декабрь']:
                        kvartal_data[date] = float(value_2.replace(',', '.').replace(' ', ''))
                    new_data[date] = float(value_1.replace(',', '.').replace(' ', ''))
                os.remove(path=path_to_docfile)

            if len(new_data) != 0:
                update_rez_file_y(new_data, kvartal_data, xlsx_path='rez_file_Y_v2.xlsx')


if __name__ == '__main__':
    main()
